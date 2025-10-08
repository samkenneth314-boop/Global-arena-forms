from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

app = Flask(__name__)
app.secret_key = "secret-key"

# Church Info
church_name = "Global Arena of Liberty Ministry"
church_address = "Kasoa – Ofaakor Branch"
church_phone = "+233 54 123 4567"
church_website = "www.globalarena.org"

# Folders
UPLOAD_FOLDER = "uploads"
TEMP_FOLDER = "temp_images"
DOCS_FOLDER = "member_docs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)
os.makedirs(DOCS_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER


@app.route('/')
def form():
    return render_template(
        "form.html",
        church_name=church_name,
        church_address=church_address,
        church_phone=church_phone,
        church_website=church_website,
        pre={}
    )

# ✅ Added this safe line (fixes your success.html redirect)
app.add_url_rule('/', 'index', form)


@app.route('/preview', methods=['POST'])
def preview():
    form_data = request.form.to_dict()
    photo = request.files.get('photo')
    temp_photo = None

    if photo and photo.filename != "":
        filename = secure_filename(photo.filename)
        temp_photo = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
        photo.save(os.path.join(TEMP_FOLDER, temp_photo))

    form_data["temp_photo"] = temp_photo
    return render_template(
        "preview.html",
        pre=form_data,
        church_name=church_name,
        church_address=church_address,
        church_phone=church_phone,
        church_website=church_website
    )


@app.route('/submit', methods=['POST'])
def submit():
    data = request.form.to_dict()
    photo = request.files.get('photo')
    temp_photo = data.get('temp_photo')

    # Use either new uploaded photo or previous temp
    if photo and photo.filename != "":
        filename = secure_filename(photo.filename)
        photo_path = os.path.join(UPLOAD_FOLDER, filename)
        photo.save(photo_path)
    elif temp_photo:
        photo_path = os.path.join(TEMP_FOLDER, temp_photo)
    else:
        photo_path = None

    first_name = data.get("first_name", "")
    last_name = data.get("last_name", "")
    full_name = f"{first_name}_{last_name}".strip()
    doc_filename = f"{full_name}.docx"
    doc_path = os.path.join(DOCS_FOLDER, doc_filename)

    document = Document()
    document.add_heading(church_name, 0)
    document.add_paragraph(f"{church_address} | {church_phone} | {church_website}")
    document.add_heading("Membership Form", level=1)
    document.add_paragraph("")

    # Photo
    if photo_path and os.path.exists(photo_path):
        document.add_picture(photo_path, width=Inches(1.5))

    # Member info
    for key, value in data.items():
        if key not in ["photo", "temp_photo", "official_name", "official_position", "official_date"]:
            document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    # Official Use Section
    document.add_paragraph("\n--- FOR OFFICIAL USE ---")
    document.add_paragraph(f"Name: {data.get('official_name', '')}")
    document.add_paragraph(f"Position: {data.get('official_position', '')}")
    document.add_paragraph(f"Date: {data.get('official_date', '')}")

    document.save(doc_path)

    # Create a simple PDF-like confirmation file
    pdf_filename = f"{full_name}.pdf"
    pdf_path = os.path.join(DOCS_FOLDER, pdf_filename)
    os.system(f'copy "{doc_path}" "{pdf_path}" >nul')

    return render_template("success.html", pdf_filename=pdf_filename, church_name=church_name)


@app.route('/download/<filename>')
def download(filename):
    return send_from_directory(DOCS_FOLDER, filename, as_attachment=True)


@app.route('/temp_images/<filename>')
def temp_images(filename):
    return send_from_directory(TEMP_FOLDER, filename)


# Admin Dashboard Routes (kept unchanged)
@app.route('/admin')
def admin_login():
    return render_template('admin-login.html')


@app.route('/admin/dashboard')
def admin_dashboard():
    files = os.listdir(DOCS_FOLDER)
    return render_template('admin-board.html', files=files)


@app.route('/admin/remove/<filename>')
def admin_remove(filename):
    path = os.path.join(DOCS_FOLDER, filename)
    if os.path.exists(path):
        os.remove(path)
        flash(f"{filename} removed successfully.")
    return redirect(url_for('admin_dashboard'))


if __name__ == "__main__":
    app.run(debug=True)
